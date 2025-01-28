import streamlit as st
from risorse.functions import get_fe_dict, connect_wc, load_files, del_res_where,weaviate_search, get_llm, full_document_audit, understood_button
from docx import Document
from time import sleep
from io import StringIO,BytesIO
from docx.shared import RGBColor 



if "api_keys" not in st.session_state:
    st.session_state.api_keys= {
        'openai_api_key': st.secrets['OpenAI_key'],
        'weaviate_URL':st.secrets['weaviate_URL'],
        'weaviate_api_key':st.secrets['weaviate_api_key']
    }

st.markdown(
    """
<style>.element-container:has(#button-after) + div button:hover {
    cursor:url("data:image/svg+xml;utf8,<svg xmlns='http://www.w3.org/2000/svg'  width='31' height='37' viewport='0 0 100 100' style='fill:black;font-size:19px;'><text y='50%'>🐈🐈</text></svg>") 16 0,auto;
 }</style>""", unsafe_allow_html=True)
st.markdown(
    """
    <style>
    
    button[kind="secondary"] {
        background: none!important;
        border: none;
        padding: 0!important;
        color: black !important;
        text-decoration: none;
        cursor: pointer;
        border: none !important;
    }
    button[kind="secondary"]:hover {
        text-decoration: none;
        color: black !important;
    }
    button[kind="secondary"]:focus {
        outline: none !important;
        box-shadow: none !important;
        color: black !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)
st.markdown(
    """<style>
                .big-font {
                    font-size:20px;
                    }
                </style>""", unsafe_allow_html=True)

if "chosen_lang" not in st.session_state:
    st.session_state.chosen_lang='en'    
st.session_state.front_end_lang=st.session_state.chosen_lang
if "llm" not in st.session_state:
    st.session_state.llm='gpt-3.5-turbo-1106'
if "audit_mode" not in st.session_state:
    st.session_state.audit_mode='Full Document'
if "docx_txt" not in st.session_state:
    st.session_state.docx_txt=''
if "audit_results" not in st.session_state:
    st.session_state.audit_results=[]

llm=get_llm(st.session_state.llm,st.session_state.api_keys)

wc_client=connect_wc(st.session_state.api_keys)

front_end_display= get_fe_dict(st.session_state.front_end_lang)

opts=st.sidebar.expander(front_end_display['setting'])

with opts:
    ce00,ce01=st.columns(2)
    with ce01:
        audit_modalities=['Full Document','Chapter']
        audit_index=audit_modalities.index(st.session_state.audit_mode)
        st.selectbox('Audit',audit_modalities,index=audit_index,key='audit_mode')
        st.markdown('<span id="button-after"></span>', unsafe_allow_html=True)
        if st.button("",type="secondary",use_container_width=True):
            st.session_state.chosen_lang='eml'

    with ce00:
        llms=['gpt-3.5-turbo-1106','gpt-4o-mini-2024-07-18']
        llm_index=llms.index(st.session_state.llm)
        st.selectbox('LLM🧠',llms,index=llm_index,key='llm')
        lang_index=['en','it','eml'].index(st.session_state.front_end_lang)
        if st.session_state.chosen_lang!='eml':
            st.selectbox(front_end_display['language'],['en','it'],index=lang_index,key='chosen_lang')
        else:
            st.selectbox(front_end_display['language'],['en','it','eml'],index=lang_index,key='chosen_lang')

resources_management=st.sidebar.expander(front_end_display['resources'])

with resources_management:
    pdf_loader=st.file_uploader(label = front_end_display['upload']['pdf'], type=["pdf"],accept_multiple_files=True)
    if pdf_loader:
        if st.button('Upload',use_container_width=True,type='primary'):
            try:
                load_files(pdf_loader,st.session_state.api_keys)
            except:
                st.warning(front_end_display['warning'])
    if wc_client.collections.exists('Docchunk'):
        ch00,ch01,ch02=st.columns([1,3,1])
        with ch01:
            front_end_display['res_list']
        weaviate_collection=wc_client.collections.get('Docchunk')
        resources_on_weaviate=list(set([item.properties['document_name'] for item in weaviate_collection.iterator()]))
        ch10,ch11,ch12=st.columns([5,1,2])
        for r in resources_on_weaviate:
            with ch10:
                st.write(r)
            with ch11:
                st.checkbox("",value=True,key=f"checkbox_{r}",label_visibility='hidden')
            with ch12:
                st.button("DEL",key=f'button_{r}',type='primary',on_click=del_res_where,args=(weaviate_collection,r,))
    else:
        resources_on_weaviate=[]
    docx_loader=st.file_uploader(label = front_end_display['upload']['word'], type=["docx"],accept_multiple_files=True)
    if docx_loader:
        glob_txt=''
        for uploaded_file in docx_loader:
            doc=Document(uploaded_file)
            doc_txt='\n'.join([p.text for p in doc.paragraphs])
            glob_txt+=doc_txt
            glob_txt+='\n'
        st.session_state.docx_txt=glob_txt

audit_results=st.sidebar.expander(front_end_display['results'])
with audit_results:
    #for res in st.session_state.audit_results:
    #    for k in res:
    #        with st.popover(k):
    #            st.write(f"Status: {k['output']}")
    #            st.write(f"Consiglio:\n {k['reason']}")
    for res in st.session_state.audit_results:
        st.write(res)
    if st.button("End Audit",type='primary',use_container_width=True):
        if st.session_state.audit_results:
            doc=Document()
            for aud in st.session_state.audit_results:
                doc.add_heading("Audit",1)
                for el in aud['chapters']:
                    for k,v in el.items():
                        chaps=', '.join(v)
                        doc.add_paragraph().add_run(f'{k}: {chaps}')
                if aud['non_conformities']:
                    doc.add_paragraph(front_end_display["noncs"])
                    r=doc.add_paragraph().add_run(', '.join(aud['non_conformities']))
                    r.font.color.rgb=RGBColor(255,0,0)
                for p in aud['result']:
                    doc.add_heading(p['point'],3)
                    r1=doc.add_paragraph()
                    r1.add_run('Status: ')
                    r2=r1.add_run(p['conformity'])
                    if p['conformity']=='NO':
                        col=RGBColor(255,0,0)
                    else:
                        col=RGBColor(0,255,0)
                    r2.font.color.rgb=col
                    doc.add_paragraph().add_run(p['reason'])
            b=BytesIO()
            doc.save(b)
            def clear():
                st.session_state.audit_result=[]
            st.download_button('Download',data=b,mime='docx',file_name="Audit.docx",type='primary',use_container_width=True,on_click=clear)





resources_in_list=[r for r in resources_on_weaviate if f"checkbox_{r}" in st.session_state]
wanted_resources=[r for r in resources_in_list if st.session_state[f"checkbox_{r}"]]

cm00,cm01=st.columns([4,6])
with cm00:
    cn1=st.container(border=True)
    with cn1:
        st.text_input(front_end_display['topic'],key='topic')
        if st.session_state.audit_mode=='Chapter':
            for r in wanted_resources:
                chaps=list(sorted(set([w.properties['chapter'] for w in wc_client.collections.get("Docchunk").iterator() if w.properties['document_name']==r])))
                st.multiselect(r,chaps,key=f'chapters_{r}',max_selections=3)
        if 'topic' in st.session_state:
            if st.session_state.topic:
                if st.session_state.audit_mode=='Full Document':
                    if st.session_state.docx_txt:
                        if st.button("Audit",key='fd_audit',type='primary',use_container_width=True):
                            with cm01:
                                with st.container(border=True):
                                    chunks=[]
                                    for r in wanted_resources:
                                        chunks+=weaviate_search(
                                            wc_client,
                                            st.session_state.topic,
                                            r
                                        )
                                    text='\n'.join(chunks)
                                    non_confs=full_document_audit(
                                            st.session_state.topic,
                                            text,
                                            st.session_state.docx_txt,
                                            llm,
                                            st.session_state.front_end_lang)
                                    for n in non_confs:
                                        st.markdown(f"""<p class="big-font">{n['point']}</p>""", unsafe_allow_html=True)
                                        colour={'YES':'green','NO':'red'}.get([n['conformity']],'green')
                                        st.markdown(f"status: :{colour}[{front_end_display['status'][n['conformity']]}]")
                                        st.write(n['reason']) 
                                        
                                    front_end_display['confirm']
                                    new_res={
                                        'chapters':[{r:['Full Document']} for r in wanted_resources],
                                        'non_conformities':[v['point'] for v in non_confs if v['conformity']=='NO'],
                                        'result':non_confs
                                    }
                                    cc100,cc101=st.columns(2)
                                    with cc100:
                                        st.button('OK',key='yb_fd',on_click=understood_button,args=(st.session_state.audit_results,new_res),type='primary',use_container_width=True)
                                    with cc101:
                                        st.button('NO',key='no_fd',type='primary',use_container_width=True)

                                
                else:
                    for r in wanted_resources:
                        if st.session_state[f'chapters_{r}']:
                            go_on=True
                        else:
                            go_on=False
                            break
                    if not st.session_state.docx_txt:
                        go_on=False
                    if go_on:
                        if st.button("Audit",key='fd_audit',type='primary',use_container_width=True):
                                with cm01:
                                    with st.container(border=True):
                                        chunks=[]
                                        for r in wanted_resources:
                                            chunks+=weaviate_search(
                                                wc_client,
                                                st.session_state.topic,
                                                r,
                                                chapters=st.session_state[f'chapters_{r}']
                                            )
                                        text='\n'.join(chunks)
                                        non_confs=full_document_audit(
                                                st.session_state.topic,
                                                text,
                                                st.session_state.docx_txt,
                                                llm,
                                                st.session_state.front_end_lang)
                                        
                                        new_res={
                                            'chapters':[{r:st.session_state[f'chapters_{r}']} for r in wanted_resources],
                                            'non_conformities':[v['point'] for v in non_confs if v['conformity']=='NO'],
                                            'result':non_confs
                                        }
                                        for n in non_confs:
                                            st.markdown(f"""<p class="big-font">{n['point']}</p>""", unsafe_allow_html=True)
                                            colour={'YES':'green','NO':'red'}.get(n['conformity'],"red")
                                            st.markdown(f"status: :{colour}[{front_end_display['status'][n['conformity']]}]")
                                            st.write(n['reason'])                                        
                                        front_end_display['confirm']
                                        cc100,cc101=st.columns(2)
                                        with cc100:
                                            st.button('OK',key='yb_cp',on_click=understood_button,args=(st.session_state.audit_results,new_res),type='primary',use_container_width=True)
                                        with cc101:
                                            st.button('NO',key='no_cp',type='primary',use_container_width=True)
                    
        with st.popover(front_end_display['help'],use_container_width=True):
            front_end_display['help_txt']







